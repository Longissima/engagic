"""Focused tests for the typed document acquisition/extraction boundary."""

import asyncio
from datetime import datetime, timedelta
import gc
import io
from pathlib import Path
import zipfile

import aiohttp
import pytest

import analysis.analyzer_async as analyzer_module
from analysis.analyzer_async import AsyncAnalyzer
from corpus.store import CorpusOriginal, sha256_hex
from parsing.pdf import PdfExtractor
from parsing.subprocess_guard import GuardCrashed
from pipeline.document_artifacts import (
    DocumentArtifact,
    DocumentFormat,
    make_artifact,
    sniff_document_format,
    verify_tls_for_url,
)


@pytest.fixture(autouse=True)
def _inline_thread_offloads_in_restricted_test_sandbox(monkeypatch):
    """Keep unit tests deterministic where the runner forbids OS threads."""

    async def inline(call, *args, **kwargs):
        await asyncio.sleep(0)
        return call(*args, **kwargs)

    monkeypatch.setattr(asyncio, "to_thread", inline)


def _ooxml_bytes(kind: str) -> bytes:
    namespace = {
        "docx": "wordprocessingml",
        "pptx": "presentationml",
        "xlsx": "spreadsheetml",
    }[kind]
    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            f"<Types><Override ContentType='{namespace}'/></Types>",
        )
    return payload.getvalue()


class _Response:
    def __init__(
        self,
        data: bytes,
        *,
        content_type: str = "application/octet-stream",
        status: int = 200,
        url: str = "",
        headers: dict | None = None,
    ):
        self.data = data
        self.status = status
        self.headers = {"Content-Type": content_type}
        self.headers.update(headers or {})
        self.url = url

    async def read(self):
        return self.data

    async def __aenter__(self):
        return self

    async def __aexit__(self, exc_type, exc, traceback):
        return False


class _Session:
    closed = False

    def __init__(self, responses):
        self.responses = list(responses)
        self.requests = []
        self.request_headers = []

    def get(self, url, *, ssl, headers=None):
        self.requests.append((url, ssl))
        self.request_headers.append(headers or {})
        response = self.responses.pop(0)
        if isinstance(response, BaseException):
            raise response
        if not response.url:
            response.url = url
        return response


class _Limiter:
    async def wait_if_needed(self, vendor):
        return None


def _configure_http(monkeypatch, analyzer, session, corpus=None):
    async def get_session():
        return session

    analyzer._get_session = get_session
    monkeypatch.setattr(analyzer_module, "get_rate_limiter", lambda: _Limiter())
    monkeypatch.setattr(analyzer_module, "get_corpus", lambda: corpus)


@pytest.mark.parametrize(
    ("data", "expected"),
    [
        (b"%PDF-1.7\n", DocumentFormat.PDF),
        (b"<!doctype html><html></html>", DocumentFormat.HTML),
        (b"\xd0\xcf\x11\xe0legacy", DocumentFormat.DOC),
        (b"{\\rtf1 meeting}", DocumentFormat.RTF),
        (_ooxml_bytes("docx"), DocumentFormat.DOCX),
        (_ooxml_bytes("pptx"), DocumentFormat.PPTX),
        (_ooxml_bytes("xlsx"), DocumentFormat.XLSX),
    ],
)
def test_media_sniffing_uses_bytes(data, expected):
    assert sniff_document_format(
        data,
        content_type="application/octet-stream",
        source_url="https://example.test/download",
    ) is expected


def test_corpus_identity_hit_avoids_network(monkeypatch):
    data = b"{\\rtf1 cached meeting minutes}"
    content_sha = sha256_hex(data)

    class Corpus:
        def __init__(self):
            self.sightings = []

        async def get_original_artifact_by_identity(self, source_url):
            return CorpusOriginal(
                data,
                content_sha,
                "application/rtf",
                last_validated_at=datetime.now(),
                last_validation_attempt_at=datetime.now(),
            )

        async def record_sighting(self, content_sha256, source_url, banana=None):
            self.sightings.append((content_sha256, source_url, banana))

        async def lookup_extraction(self, content_sha256):
            assert content_sha256 == content_sha
            return {
                "success": True,
                "text": "cached meeting minutes",
                "method": "striprtf",
                "page_count": 0,
                "ocr_pages": 0,
            }

    corpus = Corpus()
    analyzer = AsyncAnalyzer(enable_llm=False)

    async def network_must_not_run():
        raise AssertionError("corpus hit reached the network")

    analyzer._get_session = network_must_not_run
    monkeypatch.setattr(analyzer_module, "get_corpus", lambda: corpus)

    result = asyncio.run(
        analyzer.extract_document_async("https://example.test/minutes", banana="exampleCA")
    )

    assert result["text"] == "cached meeting minutes"
    assert result["content_sha256"] == content_sha
    assert result["document_format"] == "rtf"
    assert result["from_corpus"] is True
    assert corpus.sightings


def test_stale_corpus_revision_uses_conditional_validation(monkeypatch):
    data = b"%PDF-1.7 cached"
    content_sha = sha256_hex(data)

    class Corpus:
        def __init__(self):
            self.validations = []

        async def get_original_artifact_by_identity(self, source_url):
            return CorpusOriginal(
                data,
                content_sha,
                "application/pdf",
                etag='"v1"',
                last_modified="Wed, 01 Jul 2026 12:00:00 GMT",
                last_validated_at=datetime.now() - timedelta(days=2),
                last_validation_attempt_at=datetime.now() - timedelta(days=2),
            )

        async def record_validation(self, *args, **kwargs):
            self.validations.append((args, kwargs))

        async def record_validation_failure(self, *args, **kwargs):
            raise AssertionError("304 validation must not fail open")

        async def record_sighting(self, *args, **kwargs):
            return None

    corpus = Corpus()
    session = _Session(
        [
            _Response(
                b"",
                status=304,
                headers={"ETag": '"v1"'},
            )
        ]
    )
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session, corpus)

    artifact = asyncio.run(analyzer.acquire_document_async("https://example.test/a.pdf"))

    assert artifact.data == data
    assert artifact.from_corpus is True
    assert session.request_headers == [
        {
            "If-None-Match": '"v1"',
            "If-Modified-Since": "Wed, 01 Jul 2026 12:00:00 GMT",
        }
    ]
    assert len(corpus.validations) == 1


def test_rejected_validator_retries_as_unconditional_get(monkeypatch):
    old_data = b"%PDF-1.7 old"
    new_data = b"%PDF-1.7 new"
    old_sha = sha256_hex(old_data)

    class Corpus:
        def __init__(self):
            self.archives = []

        async def get_original_artifact_by_identity(self, source_url):
            return CorpusOriginal(
                old_data,
                old_sha,
                "application/pdf",
                etag='"expired"',
                last_validated_at=datetime.now() - timedelta(days=2),
                last_validation_attempt_at=datetime.now() - timedelta(days=2),
            )

        async def archive_original(self, content_sha256, **kwargs):
            self.archives.append((content_sha256, kwargs))
            return True

        async def record_validation(self, *args, **kwargs):
            return None

        async def record_validation_failure(self, *args, **kwargs):
            raise AssertionError("successful fallback must not fail open")

        async def record_sighting(self, *args, **kwargs):
            return None

    corpus = Corpus()
    session = _Session(
        [
            _Response(b"", status=412),
            _Response(
                new_data,
                content_type="application/pdf",
                headers={"ETag": '"v2"'},
            ),
        ]
    )
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session, corpus)

    artifact = asyncio.run(analyzer.acquire_document_async("https://example.test/a.pdf"))

    assert artifact.data == new_data
    assert session.request_headers == [{"If-None-Match": '"expired"'}, {}]
    assert corpus.archives[0][1]["etag"] == '"v2"'


def test_recent_revalidation_failure_serves_archive_without_network(monkeypatch):
    data = b"%PDF-1.7 durable"
    content_sha = sha256_hex(data)

    class Corpus:
        def __init__(self):
            self.sightings = 0

        async def get_original_artifact_by_identity(self, source_url):
            return CorpusOriginal(
                data,
                content_sha,
                "application/pdf",
                last_validated_at=datetime.now() - timedelta(days=2),
                last_validation_attempt_at=datetime.now() - timedelta(minutes=5),
            )

        async def record_sighting(self, *args, **kwargs):
            self.sightings += 1

    corpus = Corpus()
    analyzer = AsyncAnalyzer(enable_llm=False)

    async def network_must_not_run():
        raise AssertionError("recent validation failure reached the network")

    analyzer._get_session = network_must_not_run
    monkeypatch.setattr(analyzer_module, "get_corpus", lambda: corpus)

    artifact = asyncio.run(analyzer.acquire_document_async("https://example.test/a.pdf"))

    assert artifact.data == data
    assert artifact.from_corpus is True
    assert corpus.sightings == 1


def test_stale_corpus_revision_refreshes_changed_bytes(monkeypatch):
    old_data = b"%PDF-1.7 old"
    new_data = b"%PDF-1.7 new"
    old_sha = sha256_hex(old_data)

    class Corpus:
        def __init__(self):
            self.archives = []

        async def get_original_artifact_by_identity(self, source_url):
            return CorpusOriginal(
                old_data,
                old_sha,
                "application/pdf",
                etag='"old"',
                last_validated_at=datetime.now() - timedelta(days=2),
                last_validation_attempt_at=datetime.now() - timedelta(days=2),
            )

        async def archive_original(self, content_sha256, **kwargs):
            self.archives.append((content_sha256, kwargs))
            return True

        async def record_validation(self, *args, **kwargs):
            return None

        async def record_validation_failure(self, *args, **kwargs):
            raise AssertionError("successful refresh must not fail open")

        async def record_sighting(self, *args, **kwargs):
            return None

    corpus = Corpus()
    session = _Session(
        [
            _Response(
                new_data,
                content_type="application/pdf",
                headers={
                    "ETag": '"new"',
                    "Last-Modified": "Thu, 06 Aug 2026 12:00:00 GMT",
                },
            )
        ]
    )
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session, corpus)

    artifact = asyncio.run(analyzer.acquire_document_async("https://example.test/a.pdf"))

    assert artifact.data == new_data
    assert artifact.content_sha256 == sha256_hex(new_data)
    assert artifact.from_corpus is False
    assert corpus.archives[0][0] == sha256_hex(new_data)
    assert corpus.archives[0][1]["etag"] == '"new"'


def test_revalidation_failure_serves_archive_and_backs_off(monkeypatch):
    data = b"%PDF-1.7 durable"
    content_sha = sha256_hex(data)

    class Corpus:
        def __init__(self):
            self.failures = []

        async def get_original_artifact_by_identity(self, source_url):
            return CorpusOriginal(
                data,
                content_sha,
                "application/pdf",
                etag='"v1"',
                last_validated_at=datetime.now() - timedelta(days=2),
                last_validation_attempt_at=datetime.now() - timedelta(days=2),
            )

        async def record_validation_failure(self, *args, **kwargs):
            self.failures.append((args, kwargs))

        async def record_sighting(self, *args, **kwargs):
            return None

    corpus = Corpus()
    session = _Session([_Response(b"gone", status=404)])
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session, corpus)

    artifact = asyncio.run(analyzer.acquire_document_async("https://example.test/a.pdf"))

    assert artifact.data == data
    assert artifact.from_corpus is True
    assert len(corpus.failures) == 1


def test_concurrent_corpus_misses_share_one_download(monkeypatch):
    data = b"%PDF-1.7 one fetch"

    class SlowResponse(_Response):
        async def read(self):
            await asyncio.sleep(0.01)
            return self.data

    class Corpus:
        def __init__(self):
            self.lookups = 0
            self.archives = 0

        async def get_original_artifact_by_identity(self, source_url):
            self.lookups += 1
            return None

        async def archive_original(self, *args, **kwargs):
            self.archives += 1
            return True

        async def record_validation(self, *args, **kwargs):
            return None

        async def record_sighting(self, *args, **kwargs):
            return None

    corpus = Corpus()
    session = _Session([SlowResponse(data, content_type="application/pdf")])
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session, corpus)

    async def acquire_twice():
        return await asyncio.gather(
            analyzer.acquire_document_async("https://example.test/a.pdf"),
            analyzer.acquire_document_async("https://example.test/a.pdf"),
        )

    first, second = asyncio.run(acquire_twice())

    assert first.data == second.data == data
    assert len(session.requests) == 1
    assert corpus.lookups == 1
    assert corpus.archives == 1


def test_html_page_resolves_office_document(monkeypatch):
    page_url = "https://example.test/meeting/42"
    docx = _ooxml_bytes("docx")
    session = _Session(
        [
            _Response(
                b"<html><body><a href='/files/agenda.docx'>Agenda</a></body></html>",
                content_type="text/html",
            ),
            _Response(docx),
        ]
    )
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session)

    artifact = asyncio.run(analyzer.acquire_document_async(page_url))

    assert artifact.document_format is DocumentFormat.DOCX
    assert artifact.source_url == "https://example.test/files/agenda.docx"
    assert [request[0] for request in session.requests] == [
        page_url,
        "https://example.test/files/agenda.docx",
    ]


def test_html_without_download_yields_sanitized_text(monkeypatch):
    page = b"""
        <html><head><title>Meeting Notice</title><script>secret()</script></head>
        <body><nav>site menu</nav><form id='aspnetForm'><main><h1>City Council</h1>
        <p>The meeting begins at 6:00 PM.</p><input value='noise'></main></form></body></html>
    """
    session = _Session([_Response(page, content_type="text/html")])
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session)

    result = asyncio.run(analyzer.extract_document_async("https://example.test/notice"))

    assert result["method"] == "html_sanitized"
    assert "Meeting Notice" in result["text"]
    assert "City Council" in result["text"]
    assert "6:00 PM" in result["text"]
    assert "secret" not in result["text"]
    assert "site menu" not in result["text"]
    assert "noise" not in result["text"]


def test_office_extraction_uses_sniffed_temp_suffix(monkeypatch):
    data = _ooxml_bytes("xlsx")
    artifact = make_artifact(
        requested_url="https://example.test/download?id=5",
        source_url="https://example.test/download?id=5",
        data=data,
        content_sha256=sha256_hex(data),
    )
    analyzer = AsyncAnalyzer(enable_llm=False)
    observed = {}

    async def acquire(url, banana=None):
        return artifact

    def extract(path, *args):
        observed["suffix"] = Path(path).suffix
        return {
            "success": True,
            "text": "Budget | 100",
            "method": "openpyxl",
            "page_count": 0,
            "ocr_pages": 0,
        }

    analyzer.acquire_document_async = acquire
    monkeypatch.setattr(analyzer_module, "get_corpus", lambda: None)
    monkeypatch.setattr(analyzer_module, "_extract_pdf_in_subprocess", extract)

    result = asyncio.run(analyzer.extract_document_async(artifact.requested_url))

    assert observed["suffix"] == ".xlsx"
    assert result["document_format"] == "xlsx"


def test_extraction_releases_artifact_before_guarded_child(monkeypatch):
    data = b"%PDF-1.7 parent-buffer-release"
    content_sha = sha256_hex(data)
    analyzer = AsyncAnalyzer(enable_llm=False)

    async def acquire(url, banana=None):
        return make_artifact(
            requested_url=url,
            source_url=url,
            data=data,
            content_sha256=content_sha,
        )

    def extract(path, *args):
        del path, args
        assert not any(
            isinstance(referrer, DocumentArtifact)
            for referrer in gc.get_referrers(data)
        )
        return {
            "success": True,
            "text": "released",
            "method": "test",
            "page_count": 1,
            "ocr_pages": 0,
        }

    analyzer.acquire_document_async = acquire
    monkeypatch.setattr(analyzer_module, "get_corpus", lambda: None)
    monkeypatch.setattr(analyzer_module, "_extract_pdf_in_subprocess", extract)

    result = asyncio.run(analyzer.extract_document_async("https://example.test/a.pdf"))

    assert result["text"] == "released"
    assert result["content_sha256"] == content_sha


def test_guarded_pdf_crash_retries_without_legislative_geometry(monkeypatch):
    calls = []

    def guarded(call, args, **kwargs):
        calls.append((call, args, kwargs))
        if len(calls) == 1:
            raise GuardCrashed("native crash", exitcode=-11)
        return {
            "success": True,
            "text": "complete fallback extraction",
            "method": "pymupdf",
            "page_count": 28,
            "ocr_pages": 0,
            "ocr_pending": 0,
        }

    monkeypatch.setattr(analyzer_module, "run_guarded", guarded)

    result = analyzer_module._extract_pdf_in_subprocess(
        "/tmp/problem.pdf", 100, 150, True, 3
    )

    assert result["text"] == "complete fallback extraction"
    assert [entry[1][3] for entry in calls] == [True, False]
    assert all(entry[2]["timeout"] == 600 for entry in calls)


def test_session_rotation_waits_for_each_sessions_final_request(monkeypatch):
    class Session:
        def __init__(self, name):
            self.name = name
            self.closed = False

        async def close(self):
            self.closed = True

    analyzer = AsyncAnalyzer(enable_llm=False)
    analyzer._recycle_after = 2
    first = Session("first")
    second = Session("second")
    analyzer.http_session = first

    async def get_session():
        if analyzer.http_session is None:
            analyzer.http_session = second
        return analyzer.http_session

    analyzer._get_session = get_session

    async def rotate():
        first_lease = await analyzer._session_for_download()
        second_lease = await analyzer._session_for_download()
        assert first_lease is first
        assert second_lease is second
        assert not first.closed

        await analyzer._release_download_session(first)
        assert first.closed
        assert not second.closed

        await analyzer._release_download_session(second)
        assert not second.closed
        await analyzer.close()
        assert second.closed

    asyncio.run(rotate())


def test_docx_dispatch_uses_python_docx():
    from docx import Document

    payload = io.BytesIO()
    document = Document()
    document.add_heading("Planning Commission", level=1)
    document.add_paragraph("The hearing begins at 7:00 PM.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Case"
    table.cell(0, 1).text = "Z-42"
    document.save(payload)

    result = PdfExtractor().extract_from_bytes(payload.getvalue())

    assert result["method"] == "python-docx"
    assert "Planning Commission" in result["text"]
    assert "Case | Z-42" in result["text"]


def test_tls_verification_policy_is_narrow(monkeypatch):
    assert verify_tls_for_url("https://example.test/a.pdf") is True
    assert verify_tls_for_url("https://city.granicus.com/a.pdf") is True
    assert verify_tls_for_url("https://s3.amazonaws.com/another-bucket/a.pdf") is True
    legacy_url = "https://s3.amazonaws.com/granicus_production_attachments/a.pdf"
    assert verify_tls_for_url(legacy_url) is False
    assert verify_tls_for_url(
        "https://granicus_production_attachments.s3.amazonaws.com/a.pdf"
    ) is False

    session = _Session(
        [
            _Response(b"%PDF-1.7 normal"),
            _Response(b"%PDF-1.7 legacy"),
        ]
    )
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session)

    asyncio.run(analyzer.acquire_document_async("https://example.test/a.pdf"))
    asyncio.run(analyzer.acquire_document_async(legacy_url))

    assert [ssl for _url, ssl in session.requests] == [True, False]


def test_transient_connection_failure_retries_once(monkeypatch):
    session = _Session(
        [
            aiohttp.ClientConnectionError("reset"),
            _Response(b"%PDF-1.7 recovered", content_type="application/pdf"),
        ]
    )
    analyzer = AsyncAnalyzer(enable_llm=False)
    _configure_http(monkeypatch, analyzer, session)
    sleeps = []

    async def no_wait(attempt, retry_after):
        sleeps.append((attempt, retry_after))

    analyzer._sleep_download_retry = no_wait

    artifact = asyncio.run(
        analyzer.acquire_document_async("https://example.test/retry.pdf")
    )

    assert artifact.document_format is DocumentFormat.PDF
    assert len(session.requests) == 2
    assert sleeps == [(0, None)]
